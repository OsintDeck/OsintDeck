from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_revision_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Revisión y Actualización del TFI - OSINT Deck', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Este documento detalla los cambios sugeridos para el Trabajo Final de Licenciatura (TFI) basándose en las últimas actualizaciones del plugin OSINT Deck.')

    # Section 1
    doc.add_heading('1. Cambios Sugeridos y Justificación', level=1)
    
    changes = [
        ("Arquitectura del Sistema", 
         "Actualizar la descripción de la arquitectura de 'modular' a 'Arquitectura de Capas (Clean Architecture)'.",
         "El plugin ahora separa claramente Dominio, Infraestructura y Presentación, mejorando la escalabilidad y mantenibilidad."),
        ("Motor de Detección", 
         "Incorporar el uso de Inteligencia Artificial (Clasificador Naive Bayes).",
         "La detección ya no es solo estática por Regex, sino que aprende y detecta intenciones abstractas y lenguaje tóxico."),
        ("Gestión de Datos", 
         "Reflejar el uso de Tablas SQL Personalizadas.",
         "Mejora el rendimiento y permite manejar grandes volúmenes de herramientas sin sobrecargar WordPress."),
        ("Seguridad y Filtrado", 
         "Añadir el módulo de Detección de Toxicidad.",
         "Refuerza el carácter profesional y ético de la herramienta.")
    ]

    for item, change, justification in changes:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{item}: ")
        run.bold = True
        p.add_run(f"{change}\n")
        run_just = p.add_run("Justificación: ")
        run_just.italic = True
        p.add_run(justification)

    # Section 2
    doc.add_heading('2. Guía de Instalación', level=1)
    doc.add_paragraph('Requisitos: WordPress 6.0+, PHP 7.4+ (8.2 recomendado), MySQL 5.7+.')
    
    steps = [
        "Comprimir la carpeta 'osint-deck' en un .zip.",
        "Subir e instalar desde Plugins > Añadir nuevo en WordPress.",
        "Activar el plugin (se crearán las tablas automáticamente).",
        "Configurar ajustes básicos en el menú OSINT Deck.",
        "Insertar el shortcode [osint_deck] en cualquier página."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # Section 3
    doc.add_heading('3. Anexo: Documentación Técnica', level=1)
    doc.add_paragraph('Estructura de Carpetas:', style='List Bullet')
    doc.add_paragraph('src/Core: Inicialización.', style='List Bullet 2')
    doc.add_paragraph('src/Domain: Lógica de negocio (AI, Parser).', style='List Bullet 2')
    doc.add_paragraph('src/Infrastructure: Base de datos y logs.', style='List Bullet 2')
    doc.add_paragraph('src/Presentation: Admin y Frontend.', style='List Bullet 2')

    doc.add_paragraph('Endpoints AJAX:', style='List Bullet')
    doc.add_paragraph('osint_deck_search: Motor de búsqueda principal.', style='List Bullet 2')
    doc.add_paragraph('osint_deck_log_event: Registro de métricas.', style='List Bullet 2')

    # Section 4
    doc.add_heading('4. Conclusión de la Revisión', level=1)
    doc.add_paragraph('Las actualizaciones recientes han convertido a OSINT Deck en una plataforma mucho más robusta, ideal para un TFI de Licenciatura en Ciberseguridad, integrando conceptos de IA, seguridad defensiva y arquitectura de software moderna.')

    doc_path = r'C:\xampp\htdocs\osint.com.ar.local\wp-content\plugins\osint-deck\TFI\Nuevo_Documento_TFI_Revisado.docx'
    doc.save(doc_path)
    print(f"Documento creado en: {doc_path}")

if __name__ == "__main__":
    create_revision_doc()
