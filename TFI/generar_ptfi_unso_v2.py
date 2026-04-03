import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_margins(section):
    section.left_margin = Inches(1.18) # ~3.0 cm
    section.right_margin = Inches(0.98) # ~2.5 cm
    section.top_margin = Inches(0.98) # ~2.5 cm
    section.bottom_margin = Inches(0.98) # ~2.5 cm

def set_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_styled_paragraph(doc, text, style=None, bold=False, italic=False, align='justify', space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align == 'justify' else WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p

def generate_ptfi_unso_v2():
    doc = Document()
    
    # Configuración de secciones
    section = doc.sections[0]
    set_margins(section)

    # --- PORTADA ---
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("UNIVERSIDAD NACIONAL RAÚL SCALABRINI ORTIZ\n(UNSO)\n\nLICENCIATURA EN CIBERSEGURIDAD\n\nTRABAJO FINAL INTEGRADOR (TFI)")
    set_font(run1, size=14, bold=True)

    doc.add_paragraph("\n\n\n")
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("TÍTULO DEL PROYECTO DE TRABAJO FINAL INTEGRADOR:\n")
    set_font(run2, size=12, bold=True)
    run3 = p2.add_run("OSINT Deck: Orquestación Inteligente y Arquitectura de Capas para la Optimización del Flujo de Inteligencia de Fuentes Abiertas")
    set_font(run3, size=14, bold=True)

    doc.add_paragraph("\n\n\n\n")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run4 = p3.add_run("Autor/a: Sebastián Cendra, Claudio Pandelo, Paolo Peña Ramírez, Guillermo Quintana, Damián Radiminsky.")
    set_font(run4, size=12)

    doc.add_paragraph("\n")
    
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run5 = p4.add_run("Fecha de Presentación: Febrero de 2026")
    set_font(run5, size=12)

    doc.add_paragraph("\n\n\n")
    
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run6 = p5.add_run("CABA, Buenos Aires")
    set_font(run6, size=12)

    doc.add_page_break()

    # --- 1. RESUMEN ---
    add_styled_paragraph(doc, "1. Resumen del Proyecto", bold=True)
    resumen = (
        "El presente Proyecto de Trabajo Final Integrador (PTFI) propone el desarrollo de OSINT Deck, una plataforma de orquestación "
        "diseñada para centralizar y automatizar el flujo de inteligencia de fuentes abiertas (OSINT). El proyecto aborda la problemática "
        "de la carga cognitiva y la fragmentación operativa en investigaciones digitales mediante una arquitectura de capas (Clean Architecture) "
        "y un motor de decisión híbrido que integra Inteligencia Artificial (Naive Bayes). Se busca demostrar cómo la automatización del triage "
        "de indicadores de compromiso (IOCs) y la aplicación de estándares OWASP 2025 optimizan la respuesta ante incidentes. El resultado "
        "esperado es una herramienta profesional que trasciende el modelo de directorio estático, ofreciendo una solución escalable y segura "
        "alineada con las competencias de la Licenciatura en Ciberseguridad."
    )
    add_styled_paragraph(doc, resumen)

    # --- 2. VINCULACIÓN CON EL PERFIL PROFESIONAL (NUEVO) ---
    add_styled_paragraph(doc, "2. Vinculación con el Perfil Profesional de la Licenciatura", bold=True)
    vinculacion = (
        "El proyecto OSINT Deck constituye una síntesis integradora de las competencias adquiridas durante la Licenciatura en Ciberseguridad de la UNSO. "
        "Se correlaciona directamente con los objetivos de la carrera en tres ejes fundamentales: \n"
        "1) Ciberinteligencia y Vigilancia Tecnológica: Al proponer un sistema de recolección y análisis proactivo de información pública para la detección de amenazas. \n"
        "2) Desarrollo Seguro de Software (Security by Design): Al implementar una arquitectura de capas auditada bajo estándares internacionales (OWASP 2025), "
        "demostrando la capacidad de diseñar sistemas resilientes. \n"
        "3) Respuesta ante Incidentes: Al optimizar los tiempos de respuesta mediante la orquestación automática de recursos, facilitando el análisis forense "
        "preliminar y la mitigación de ataques basados en ingeniería social o filtración de datos."
    )
    add_styled_paragraph(doc, vinculacion)

    # --- 3. PLANTEAMIENTO DEL PROBLEMA Y ESTADO DEL ARTE ---
    add_styled_paragraph(doc, "3. Planteamiento del Problema y Estado de la Técnica", bold=True)
    problema = (
        "En la actualidad, el ecosistema OSINT se caracteriza por una paradoja de sobreabundancia informativa. Frameworks como 'OSINT Framework' "
        "u 'Awesome OSINT' funcionan como catálogos estáticos de enlaces que, si bien son valiosos, obligan al analista a una verificación manual "
        "constante y a una carga cognitiva elevada (Heuer, 1999). Herramientas más avanzadas como 'Maltego' ofrecen capacidades de correlación, "
        "pero su curva de aprendizaje y costos limitan su accesibilidad. OSINT Deck se sitúa en un punto intermedio crítico: ofrece la simplicidad "
        "de una interfaz web con la potencia de un motor de decisión inteligente. El problema central radica en la ausencia de herramientas que "
        "armonicen la automatización del triage con un entorno de ejecución seguro y accesible, protegiendo tanto al investigador como a la "
        "integridad de la investigación."
    )
    add_styled_paragraph(doc, problema)

    # --- 4. OBJETIVOS ---
    add_styled_paragraph(doc, "4. Objetivos", bold=True)
    add_styled_paragraph(doc, "4.1 Objetivo General", italic=True)
    add_styled_paragraph(doc, "Desarrollar una plataforma orquestadora de servicios OSINT basada en una librería core propietaria y una arquitectura de capas, orientada a la optimización de procesos de ciberinteligencia y respuesta ante incidentes.")
    
    add_styled_paragraph(doc, "4.2 Objetivos Específicos", italic=True)
    objs = [
        "Desarrollar la 'OSINT Core Library' bajo estándares PSR-4 para garantizar un desacoplamiento total de la lógica de negocio respecto al CMS.",
        "Implementar un motor de clasificación probabilístico basado en Naive Bayes para la detección de intenciones en lenguaje natural.",
        "Diseñar un sistema de persistencia híbrido utilizando tablas SQL personalizadas para garantizar la integridad y trazabilidad de los datos (Logs/Herramientas).",
        "Auditar la plataforma bajo el estándar OWASP Top 10:2025, enfocándose en la mitigación de inyecciones y fallas en el control de acceso.",
        "Establecer un marco ético-legal basado en la Ley 25.326 para la recolección responsable de datos en fuentes abiertas."
    ]
    for obj in objs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(obj)
        set_font(run)

    # --- 5. MARCO TEÓRICO (EXPANDIDO) ---
    add_styled_paragraph(doc, "5. Marco Teórico", bold=True)
    marco = (
        "El sustento teórico de este trabajo trasciende la mera descripción de herramientas. Se apoya en la 'Teoría de la Inteligencia Centrada en el Objetivo' "
        "(Clark, 2019), que propone que la organización de la información debe responder al problema de análisis y no a la herramienta. "
        "Desde la ciberseguridad, se adopta el modelo de 'Defensa en Profundidad' aplicado al desarrollo de aplicaciones web, integrando controles "
        "técnicos (Nonces, Sanitización) y arquitectónicos (Separación de Capas). El uso de IA mediante el Teorema de Bayes se fundamenta como una "
        "solución técnica a las limitaciones del análisis heurístico tradicional, permitiendo una interpretación contextual de los indicadores de "
        "compromiso (IOCs) analizados."
    )
    add_styled_paragraph(doc, marco)

    # --- 6. METODOLOGÍA ---
    add_styled_paragraph(doc, "6. Metodología y Técnicas", bold=True)
    metodologia = (
        "Se adoptará un enfoque de investigación aplicada mediante el método de desarrollo iterativo e incremental. Se priorizará la "
        "seguridad por diseño (Security by Design) realizando auditorías continuas en cada sprint. Las técnicas de recolección de datos "
        "incluyen el análisis de patrones de dorks y la normalización de inputs técnicos. La validación se llevará a cabo mediante "
        "pruebas de regresión y auditorías de caja blanca sobre la librería de dominio propia, asegurando que la lógica de decisión sea "
        "consistente y resistente a manipulaciones externas."
    )
    add_styled_paragraph(doc, metodologia)

    # --- 7. CRONOGRAMA ---
    add_styled_paragraph(doc, "7. Cronograma", bold=True)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fase'
    hdr_cells[1].text = 'Actividad'
    hdr_cells[2].text = 'Semanas'
    
    data = [
        ("Fase 1: Planificación", "Revisión bibliográfica y mapeo de competencias UNSO", "1-3"),
        ("Fase 1: Planificación", "Redacción del plan de trabajo y diseño de arquitectura", "4-5"),
        ("Fase 2: Ejecución", "Desarrollo de la OSINT Core Library (Domain Layer)", "6-9"),
        ("Fase 2: Ejecución", "Integración con WordPress y motor de IA", "10-12"),
        ("Fase 2: Ejecución", "Auditoría OWASP 2025 y QA de seguridad", "13-14"),
        ("Fase 3: Cierre", "Análisis de impacto y redacción del TFI final", "15-16")
    ]
    
    for fase, act, sem in data:
        row_cells = table.add_row().cells
        row_cells[0].text = fase
        row_cells[1].text = act
        row_cells[2].text = sem

    doc.add_page_break()

    # --- 8. REFERENCIAS (PROFESIONALIZADAS) ---
    add_styled_paragraph(doc, "8. Referencias Bibliográficas", bold=True)
    refs = [
        "Clark, R. M. (2019). Intelligence Analysis: A Target-Centric Approach. CQ Press.",
        "Heuer, R. J. (1999). Psychology of Intelligence Analysis. Center for the Study of Intelligence.",
        "Kent, S. (1949). Strategic Intelligence for American World Policy. Princeton University Press.",
        "OWASP. (2025). Top 10 Web Application Security Risks. https://owasp.org/Top10/2025/",
        "Universidad Nacional Raúl Scalabrini Ortiz. (2019). Resolución Ministerial N° 4313: Licenciatura en Ciberseguridad."
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run()
        set_font(run)

    # Guardar
    path = r'C:\xampp\htdocs\osint.com.ar.local\wp-content\plugins\osint-deck\TFI\PTFI_Osint_Deck_UNSO_v2.docx'
    doc.save(path)
    print(f"Proyecto de TFI (PTFI) v2 generado en: {path}")

if __name__ == "__main__":
    generate_ptfi_unso_v2()
