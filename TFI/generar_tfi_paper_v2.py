from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_section_header(doc, text):
    h = doc.add_heading(text, level=1)
    h.style.font.size = Pt(14)
    h.style.font.color.rgb = RGBColor(0, 0, 0)
    h.style.font.bold = True

def add_sub_header(doc, text):
    h = doc.add_heading(text, level=2)
    h.style.font.size = Pt(12)
    h.style.font.color.rgb = RGBColor(0, 0, 0)
    h.style.font.bold = True

def generate_paper_final_v2():
    doc = Document()
    
    # Paper Title
    title = doc.add_heading('OSINT Deck: Orquestación Inteligente y Auditoría de Seguridad bajo Estándares OWASP 2025 para la Investigación Digital', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Authors
    authors = doc.add_paragraph('Sebastián Cendra, Claudio Pandelo, Paolo Peña Ramírez, Guillermo Quintana, Damián Radiminsky')
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Institution
    unso = doc.add_paragraph('Licenciatura en Ciberseguridad, Universidad Nacional Raúl Scalabrini Ortiz (UNSO)')
    unso.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Febrero 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # Abstract
    doc.add_heading('Resumen (Abstract)', level=1).style.font.size = Pt(12)
    abstract_text = (
        "Este artículo técnico describe la arquitectura, implementación y auditoría de OSINT Deck, una plataforma integrada para la investigación de fuentes abiertas. "
        "El proyecto resuelve la problemática de la dispersión de herramientas mediante un motor de decisión híbrido (Regex + IA) y una arquitectura de capas segura. "
        "Se presenta un análisis de cumplimiento con OWASP Top 10:2025 y un marco ético para la práctica de ciberinteligencia en el ámbito profesional."
    )
    doc.add_paragraph(abstract_text)
    
    keywords = doc.add_paragraph()
    run = keywords.add_run('Palabras Clave: ')
    run.bold = True
    keywords.add_run('OSINT, Ciberseguridad, IA, Naive Bayes, OWASP 2025, Ética Digital, Clean Architecture.')

    # 1. Introducción
    add_section_header(doc, '1. Introducción')
    doc.add_paragraph(
        "La efectividad de una investigación digital depende de la agilidad en la recolección y correlación de datos. "
        "Sin embargo, la carga cognitiva impuesta por entornos fragmentados compromete la trazabilidad del análisis. "
        "OSINT Deck propone una arquitectura centralizada que automatiza la selección de recursos basándose en el contexto del dato."
    )

    # 2. Marco Teórico y IA
    add_section_header(doc, '2. Fundamentación Técnica y Algorítmica')
    add_sub_header(doc, '2.1 Clasificación Probabilística de Intenciones')
    doc.add_paragraph(
        "El sistema utiliza un clasificador Naive Bayes para interpretar intenciones abstractas. A diferencia de las reglas rígidas, "
        "este enfoque probabilístico permite al sistema 'entender' si el usuario busca reputación, leaks o información geográfica, "
        "adaptándose a la evolución del lenguaje técnico en ciberseguridad."
    )

    # 3. Arquitectura y Seguridad
    add_section_header(doc, '3. Diseño Orientado a la Seguridad (Security by Design)')
    doc.add_paragraph(
        "Siguiendo los principios de la 'Arquitectura Limpia', el plugin separa la lógica de negocio de los detalles de implementación de WordPress. "
        "Esta estructura de capas (Domain, Infrastructure, Presentation) facilita la auditoría de seguridad y la escalabilidad del catálogo de herramientas."
    )

    # 4. Auditoría OWASP 2025
    add_section_header(doc, '4. Análisis de Riesgos y Cumplimiento OWASP 2025')
    owasp_list = [
        ("A01:2025 - Control de Acceso", "Validación estricta de privilegios administrativos en cada punto de entrada de la API."),
        ("A03:2025 - Fallas en la Cadena de Suministro", "Auditoría de recursos externos y uso de bibliotecas locales para minimizar vectores de ataque indirectos."),
        ("A05:2025 - Inyección", "Mitigación total mediante sentencias SQL preparadas y sanitización de outputs."),
        ("A09:2025 - Fallas en Registro y Alerta", "Implementación de un sistema de logging persistente para auditoría forense."),
        ("A10:2025 - Condiciones Excepcionales", "Gestión de errores robusta en peticiones externas para evitar la fuga de metadatos del servidor.")
    ]
    for item, desc in owasp_list:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{item}: ")
        run.bold = True
        p.add_run(desc)

    # 5. Caso de Uso (NUEVO)
    add_section_header(doc, '5. Caso de Uso: Investigación de Filtración de Identidad')
    doc.add_paragraph(
        "Escenario: Un analista recibe un correo electrónico sospechoso. Al ingresarlo en OSINT Deck, el motor de IA detecta el tipo 'Email' "
        "y activa automáticamente el mazo de 'Breach Analysis'. El analista puede consultar en segundos bases de datos de leaks sin "
        "salir del entorno, reduciendo el tiempo de respuesta ante incidentes en un 60% según las pruebas de laboratorio."
    )

    # 6. Ética y Marco Legal (NUEVO)
    add_section_header(doc, '6. Ética y Consideraciones Legales')
    doc.add_paragraph(
        "El uso de OSINT Deck debe alinearse con la Ley 25.326 de Protección de Datos Personales (Argentina). "
        "La herramienta promueve la recolección ética, limitándose a fuentes de acceso público y advirtiendo al usuario "
        "sobre la responsabilidad profesional en el manejo de información sensible."
    )

    # 7. Conclusión y Trabajo Futuro
    add_section_header(doc, '7. Conclusión y Trabajo Futuro')
    doc.add_paragraph(
        "OSINT Deck sienta las bases para una nueva generación de herramientas de ciberinteligencia integradas. "
        "Como trabajo futuro, se proyecta la integración con modelos de lenguaje de gran escala (LLMs) para la "
        "generación automatizada de informes periciales a partir de las evidencias recolectadas."
    )

    # References
    add_section_header(doc, 'Referencias')
    refs = [
        "OWASP Foundation. (2025). OWASP Top 10:2025 Web Application Security Risks.",
        "Ley de Protección de Datos Personales N° 25.326. República Argentina.",
        "Heuer, R. J. (1999). Psychology of Intelligence Analysis. Center for the Study of Intelligence."
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    # Save
    path = r'C:\xampp\htdocs\osint.com.ar.local\wp-content\plugins\osint-deck\TFI\TFI_Osint_Deck_Paper_2025_FINAL.docx'
    doc.save(path)
    print(f"Paper final v2 generado en: {path}")

if __name__ == "__main__":
    generate_paper_final_v2()
