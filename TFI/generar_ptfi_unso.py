import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_margins(section):
    section.left_margin = Inches(1.18) # ~3.0 cm
    section.right_margin = Inches(0.98) # ~2.5 cm
    section.top_margin = Inches(0.98) # ~2.5 cm
    section.bottom_margin = Inches(0.98) # ~2.5 cm

def set_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_styled_paragraph(doc, text, style=None, bold=False, italic=False, align='justify', space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align == 'justify' else WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p

def generate_ptfi_unso():
    doc = Document()
    
    # Configuración de secciones
    section = doc.sections[0]
    set_margins(section)

    # --- PORTADA ---
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("UNIVERSIDAD NACIONAL RAÚL SCALABRINI ORTIZ\n(UNSO)\n\nLICENCIATURA EN CIBERSEGURIDAD\n\nTRABAJO FINAL INTEGRADOR (TFI)")
    set_font(run1, size=14, bold=True)

    doc.add_paragraph("\n\n\n\n")
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("TÍTULO DEL PROYECTO DE TRABAJO FINAL INTEGRADOR:\n")
    set_font(run2, size=12, bold=True)
    run3 = p2.add_run("OSINT Deck: Desarrollo de una plataforma inteligente de orquestación de herramientas para investigaciones digitales bajo estándares OWASP 2025")
    set_font(run3, size=14, bold=True)

    doc.add_paragraph("\n\n\n\n")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run4 = p3.add_run("Autor/a: Sebastián Cendra, Claudio Pandelo, Paolo Peña Ramírez, Guillermo Quintana, Damián Radiminsky.")
    set_font(run4, size=12)

    doc.add_paragraph("\n")
    
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run5 = p4.add_run("Fecha de Presentación: Febrero de 2026")
    set_font(run5, size=12)

    doc.add_paragraph("\n\n\n")
    
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run6 = p5.add_run("CABA, Buenos Aires")
    set_font(run6, size=12)

    doc.add_page_break()

    # --- ÍNDICE (Marcador de posición) ---
    add_styled_paragraph(doc, "ÍNDICE GENERAL", bold=True, align='center')
    doc.add_paragraph("(El índice se generará automáticamente en la versión final del documento Word)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 1. RESUMEN ---
    add_styled_paragraph(doc, "1. Resumen del Proyecto", bold=True)
    resumen = (
        "El presente Proyecto de Trabajo Final Integrador (PTFI) propone el desarrollo de OSINT Deck, un plugin para WordPress "
        "diseñado para centralizar y automatizar el acceso a herramientas de investigación digital de fuentes abiertas (OSINT). "
        "El problema abordado es la fragmentación y carga cognitiva que enfrentan los analistas al utilizar múltiples plataformas dispersas. "
        "La metodología combina el desarrollo ágil con una arquitectura de capas (Clean Architecture) y un motor de decisión híbrido "
        "basado en Inteligencia Artificial (Naive Bayes) para la clasificación de intenciones de búsqueda. "
        "Como resultado esperado, se entregará una herramienta funcional auditada bajo el estándar OWASP Top 10:2025, mejorando "
        "la eficiencia operativa en investigaciones de ciberseguridad en un entorno web seguro."
    )
    add_styled_paragraph(doc, resumen)

    # --- 2. JUSTIFICACIÓN ---
    add_styled_paragraph(doc, "2. Justificación", bold=True)
    justificacion = (
        "La ciberinteligencia moderna requiere agilidad y precisión. La dispersión actual de recursos OSINT genera tiempos muertos "
        "y aumenta el riesgo de omitir fuentes críticas. OSINT Deck es relevante porque transforma catálogos estáticos en un entorno "
        "operativo dinámico. Su importancia radica en la integración de capacidades de detección inteligente y la mitigación de "
        "vulnerabilidades web según estándares internacionales (OWASP 2025). Este proyecto aporta a la comunidad de ciberseguridad "
        "una herramienta de código abierto que estandariza y asegura el flujo de trabajo investigativo, alineándose con las competencias "
        "de la Licenciatura en Ciberseguridad de la UNSO."
    )
    add_styled_paragraph(doc, justificacion)

    # --- 3. PLANTEAMIENTO DEL PROBLEMA ---
    add_styled_paragraph(doc, "3. Planteamiento del tema/problema", bold=True)
    problema = (
        "En la práctica de la investigación digital, el analista debe navegar manualmente entre decenas de plataformas (Shodan, VirusTotal, "
        "IntelTechniques, etc.), lo que fragmenta el flujo de análisis. Esta situación se agrava por la obsolescencia de los directorios "
        "OSINT actuales, donde muchos enlaces están inactivos. El problema central se define como la ineficiencia operativa y la falta "
        "de un entorno seguro y unificado que asista al investigador mediante la detección automática del contexto del dato (IP, dominio, "
        "email, etc.). ¿Cómo puede una arquitectura de software inteligente y segura optimizar la orquestación de herramientas OSINT y "
        "reducir la superficie de ataque en aplicaciones de investigación basadas en CMS?"
    )
    add_styled_paragraph(doc, problema)

    # --- 4. OBJETIVOS ---
    add_styled_paragraph(doc, "4. Objetivos", bold=True)
    add_styled_paragraph(doc, "4.1 Objetivo General", italic=True)
    add_styled_paragraph(doc, "Desarrollar un plugin para WordPress basado en una arquitectura de capas segura que integre y automatice el uso de herramientas OSINT mediante inteligencia artificial.")
    
    add_styled_paragraph(doc, "4.2 Objetivos Específicos", italic=True)
    objs = [
        "Desarrollar una librería de dominio propietaria (OSINT Core Library) para el procesamiento y clasificación de datos de investigación.",
        "Diseñar una arquitectura de capas (Domain-Driven Design) que aísle la lógica de investigación de la infraestructura web.",
        "Implementar un motor de clasificación probabilístico (Naive Bayes) dentro de la librería core para identificar intenciones de búsqueda.",
        "Desarrollar un sistema de gestión de persistencia mediante tablas SQL personalizadas para optimizar el rendimiento.",
        "Auditar y mitigar riesgos de seguridad siguiendo el estándar OWASP Top 10 versión 2025.",
        "Validar la efectividad de la herramienta mediante casos de uso reales de investigación digital."
    ]
    for obj in objs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(obj)
        set_font(run)

    # --- 5. HIPÓTESIS / PREGUNTA DE INVESTIGACIÓN ---
    add_styled_paragraph(doc, "5. Hipótesis / Pregunta de Investigación", bold=True)
    hipotesis = (
        "La implementación de un motor de decisión híbrido basado en Inteligencia Artificial y una arquitectura de capas segura permite "
        "reducir el tiempo de recolección de información en un 40% y mitigar las vulnerabilidades críticas del estándar OWASP 2025 en "
        "comparación con el uso de directorios OSINT tradicionales."
    )
    add_styled_paragraph(doc, hipotesis)

    # --- 6. MARCO TEÓRICO ---
    add_styled_paragraph(doc, "6. Marco Teórico (preliminar)", bold=True)
    marco = (
        "El marco teórico se sustenta en la Psicología del Análisis de Inteligencia (Heuer, 1999) y el Ciclo de Inteligencia (Kent, 1949). "
        "Técnicamente, se fundamenta en el Teorema de Bayes para la clasificación de texto y en los principios de Clean Architecture para "
        "el diseño de software seguro. Asimismo, se incorpora el estándar OWASP Top 10:2025 como marco de referencia para la seguridad "
        "de aplicaciones web, enfocándose en la mitigación de fallas en el control de acceso (A01), inyecciones (A05) y fallas en la cadena "
        "de suministro (A03). Se considera además la Ley 25.326 de Protección de Datos Personales en Argentina como base legal para la "
        "recolección ética de información de fuentes abiertas."
    )
    add_styled_paragraph(doc, marco)

    # --- 7. METODOLOGÍA ---
    add_styled_paragraph(doc, "7. Metodología y técnicas a utilizar", bold=True)
    metodologia = (
        "Se utilizará una metodología de desarrollo ágil (Scrum) centrada en la seguridad (Security by Design). Las técnicas incluyen: "
        "1) Desarrollo de una librería de dominio propia y desacoplada (OSINT Core Library) bajo estándares PSR-4, lo que permite que la lógica "
        "de clasificación y decisión sea independiente del CMS. 2) Análisis comparativo de herramientas OSINT. 3) Modelado de amenazas basado "
        "en el flujo de datos del plugin. 4) Desarrollo en PHP y JavaScript respetando los hooks de WordPress. 5) Entrenamiento del clasificador "
        "Naive Bayes con datasets de dorks OSINT. 6) Pruebas de seguridad estáticas y dinámicas para validar el cumplimiento con OWASP 2025."
    )
    add_styled_paragraph(doc, metodologia)

    # --- 8. CRONOGRAMA ---
    add_styled_paragraph(doc, "8. Cronograma", bold=True)
    doc.add_paragraph("A continuación se detalla la planificación temporal de las actividades del proyecto:")
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fase'
    hdr_cells[1].text = 'Actividad'
    hdr_cells[2].text = 'Semanas'
    
    data = [
        ("Fase 1: Planificación", "Revisión bibliográfica y delimitación del tema", "1-3"),
        ("Fase 1: Planificación", "Redacción del plan de trabajo", "4-5"),
        ("Fase 2: Ejecución", "Diseño de arquitectura y metodología", "6-7"),
        ("Fase 2: Ejecución", "Desarrollo del motor de IA y Core", "8-12"),
        ("Fase 2: Ejecución", "Auditoría de seguridad OWASP y Pruebas", "13-14"),
        ("Fase 3: Cierre", "Análisis de resultados y redacción final", "15-16")
    ]
    
    for fase, act, sem in data:
        row_cells = table.add_row().cells
        row_cells[0].text = fase
        row_cells[1].text = act
        row_cells[2].text = sem

    doc.add_paragraph("\n")

    # --- 9. REFERENCIAS ---
    add_styled_paragraph(doc, "9. Referencias bibliográficas y bibliografía (preliminar)", bold=True)
    refs = [
        "Fernández, J. (2023). La ciberseguridad como prioridad empresarial. Revista de Ciberinteligencia.",
        "Heuer, R. J. (1999). Psychology of Intelligence Analysis. Center for the Study of Intelligence.",
        "Kent, S. (1949). Strategic Intelligence for American World Policy. Princeton University Press.",
        "OWASP. (2025). Top 10 Web Application Security Risks. https://owasp.org/Top10/2025/",
        "Universidad Nacional Raúl Scalabrini Ortiz. (2024). Reglamento de Trabajo Final de Licenciatura."
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run()
        set_font(run)

    # --- 10. ANEXOS ---
    add_styled_paragraph(doc, "10. Anexos", bold=True)
    doc.add_paragraph("Anexo A: Código fuente del Clasificador Naive Bayes (src/Domain/Service/NaiveBayesClassifier.php)")
    doc.add_paragraph("Anexo B: Mapeo detallado de controles OWASP 2025 implementados.")

    # Guardar
    path = r'C:\xampp\htdocs\osint.com.ar.local\wp-content\plugins\osint-deck\TFI\PTFI_Osint_Deck_UNSO.docx'
    doc.save(path)
    print(f"Proyecto de TFI (PTFI) generado en: {path}")

if __name__ == "__main__":
    generate_ptfi_unso()
