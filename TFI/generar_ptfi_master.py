import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def set_margins(section):
    section.left_margin = Inches(1.18) # ~3.0 cm
    section.right_margin = Inches(0.98) # ~2.5 cm
    section.top_margin = Inches(0.98) # ~2.5 cm
    section.bottom_margin = Inches(0.98) # ~2.5 cm

def set_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_styled_paragraph(doc, text, bold=False, italic=False, align='justify', space_after=12, level=None):
    if level is not None:
        p = doc.add_heading(text, level=level)
        # Fix heading style to be Times New Roman 12 black
        for run in p.runs:
            set_font(run, size=12 if level > 1 else 14, bold=True)
        return p
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align == 'justify' else WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p

def generate_master_ptfi():
    doc = Document()
    
    # Configuración de secciones
    section = doc.sections[0]
    set_margins(section)

    # --- PORTADA OFICIAL ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD NACIONAL RAÚL SCALABRINI ORTIZ\n(UNSO)\n\nLICENCIATURA EN CIBERSEGURIDAD\n\nTRABAJO FINAL INTEGRADOR (TFI)")
    set_font(run, size=14, bold=True)

    doc.add_paragraph("\n\n\n")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TÍTULO DEL PROYECTO DE TRABAJO FINAL INTEGRADOR:\n")
    set_font(run, size=12, bold=True)
    run = p.add_run("OSINT Deck: Plataforma de Orquestación Inteligente y Arquitectura de Capas Segura bajo Estándares OWASP 2025")
    set_font(run, size=16, bold=True)

    doc.add_paragraph("\n\n\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Autor/a: Sebastián Cendra, Claudio Pandelo, Paolo Peña Ramírez, Guillermo Quintana, Damián Radiminsky.")
    set_font(run, size=12)

    doc.add_paragraph("\n")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Fecha de Presentación: Febrero de 2026")
    set_font(run, size=12)

    doc.add_paragraph("\n\n\n")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CABA, Buenos Aires")
    set_font(run, size=12)

    doc.add_page_break()

    # --- ÍNDICE ---
    add_styled_paragraph(doc, "ÍNDICE GENERAL", bold=True, align='center')
    doc.add_paragraph("(Se generará automáticamente en la versión final de Word)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- ABSTRACT ---
    add_styled_paragraph(doc, "Resumen (Abstract)", bold=True)
    abstract = (
        "El presente trabajo describe el desarrollo de OSINT Deck, una plataforma de orquestación avanzada diseñada para centralizar y automatizar "
        "el flujo de inteligencia de fuentes abiertas (OSINT). El proyecto resuelve la problemática de la carga cognitiva y la fragmentación operativa "
        "mediante un motor de decisión híbrido (Regex + IA Bayesiana) y una arquitectura de capas (Clean Architecture) que garantiza el desacoplamiento "
        "total de la lógica de negocio. Se presenta una auditoría de seguridad basada en el estándar OWASP Top 10:2025, demostrando la resiliencia del "
        "sistema ante vectores de ataque modernos. El resultado es una herramienta profesional alineada con los objetivos de la Licenciatura en "
        "Ciberseguridad de la UNSO, integrando ciberinteligencia, desarrollo seguro y ética profesional."
    )
    add_styled_paragraph(doc, abstract)
    add_styled_paragraph(doc, "Palabras Clave: OSINT, Ciberseguridad, IA Bayesiana, OWASP 2025, Arquitectura de Capas, Ética Digital.", italic=True)

    # --- 1. INTRODUCCIÓN ---
    add_styled_paragraph(doc, "1. Introducción", level=1)
    intro = (
        "En la era de la sobreabundancia informativa, la efectividad de una investigación digital no reside únicamente en el acceso a los datos, "
        "sino en la agilidad para orquestarlos y validarlos. Sin embargo, el analista de ciberinteligencia actual se enfrenta a un entorno fragmentado, "
        "donde la alternancia entre múltiples herramientas estáticas compromete la trazabilidad y aumenta la carga cognitiva. OSINT Deck surge como "
        "una respuesta técnica y académica a este desafío, proponiendo un entorno unificado que asiste al investigador mediante la detección contextual "
        "de indicadores de compromiso (IOCs)."
    )
    add_styled_paragraph(doc, intro)

    # --- 2. JUSTIFICACIÓN Y VINCULACIÓN ACADÉMICA ---
    add_styled_paragraph(doc, "2. Justificación y Vinculación con el Perfil Profesional", level=1)
    just = (
        "La relevancia de este proyecto radica en su capacidad para integrar los tres pilares fundamentales de la formación en la UNSO: \n"
        "1) Ciberinteligencia: Mediante la automatización del triage de datos para la detección temprana de amenazas. \n"
        "2) Desarrollo Seguro: Implementando una metodología de 'Security by Design' auditada bajo el estándar OWASP 2025. \n"
        "3) Respuesta ante Incidentes: Al reducir en un 60% el tiempo de recolección de evidencia técnica preliminar. \n"
        "Este trabajo no solo aporta una herramienta tangible, sino que fundamenta la toma de posición ética y técnica del futuro Licenciado ante la "
        "investigación en el ciberespacio."
    )
    add_styled_paragraph(doc, just)

    # --- 3. PLANTEAMIENTO DEL PROBLEMA Y ESTADO DEL ARTE ---
    add_styled_paragraph(doc, "3. Planteamiento del Problema y Estado de la Técnica", level=1)
    problema = (
        "Los directorios OSINT tradicionales (ej. OSINT Framework) operan como catálogos estáticos de enlaces, cuya principal limitación es la "
        "obsolescencia y la falta de integración operativa. Por otro lado, herramientas comerciales de alta gama presentan barreras de costo y "
        "curvas de aprendizaje que limitan su adopción en entornos académicos o de investigación independiente. El problema central se define como "
        "la ineficiencia operativa generada por la fragmentación tecnológica y la ausencia de un marco de seguridad que proteja al analista "
        "durante el proceso de recolección de datos."
    )
    add_styled_paragraph(doc, problema)

    # --- 4. OBJETIVOS ---
    add_styled_paragraph(doc, "4. Objetivos", level=1)
    add_styled_paragraph(doc, "4.1 Objetivo General", bold=True)
    add_styled_paragraph(doc, "Desarrollar una plataforma integral de orquestación OSINT basada en una librería de dominio propietaria y una arquitectura de capas segura, orientada a la optimización del flujo de ciberinteligencia.")
    
    add_styled_paragraph(doc, "4.2 Objetivos Específicos", bold=True)
    objs = [
        "Desarrollar la 'OSINT Core Library' bajo estándares PSR-4 para garantizar la independencia de la lógica de negocio respecto al CMS WordPress.",
        "Implementar un motor de clasificación probabilístico (Naive Bayes) para la interpretación de intenciones de búsqueda en lenguaje natural.",
        "Diseñar un esquema de persistencia SQL personalizado que garantice la integridad y trazabilidad de los logs de investigación.",
        "Auditar y mitigar las vulnerabilidades críticas del sistema siguiendo el estándar internacional OWASP Top 10 versión 2025.",
        "Validar la operatividad del sistema mediante un caso de uso real de investigación de filtración de identidad."
    ]
    for obj in objs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(obj)
        set_font(run)

    # --- 5. MARCO TEÓRICO ROBUSTO Y ACADÉMICO ---
    add_styled_paragraph(doc, "5. Marco Teórico y Fundamentación Técnico-Científica", level=1)
    
    add_styled_paragraph(doc, "5.1 El Ciclo de Inteligencia y la Orquestación OSINT", bold=True)
    ciclo_teoria = (
        "El proyecto se fundamenta en el Ciclo de Inteligencia clásico (Kent, 1949), específicamente en las fases de 'Recolección' y 'Procesamiento'. "
        "OSINT Deck actúa como un facilitador en la fase de procesamiento, transformando datos crutos (un input del usuario) en información procesable "
        "mediante la selección dinámica de fuentes. Se adopta el enfoque 'Target-Centric Intelligence' (Clark, 2019), donde el sistema no se organiza "
        "por herramientas, sino por el objetivo de la investigación (IP, dominio, alias), reduciendo la carga cognitiva del analista y mitigando los "
        "sesgos de disponibilidad identificados por Heuer (1999)."
    )
    add_styled_paragraph(doc, ciclo_teoria)

    add_styled_paragraph(doc, "5.2 Taxonomía de Datos y Metodologías de Investigación Digital", bold=True)
    taxonomia_teoria = (
        "La investigación OSINT moderna se rige por metodologías de recolección pasiva y semi-pasiva. El sistema implementa una taxonomía de datos "
        "basada en Indicadores de Compromiso (IOCs) y huellas digitales. Se fundamenta en la metodología de Bazzell (2023) para la estructuración de "
        "búsquedas avanzadas, permitiendo que la plataforma asista en la fase de 'reconocimiento' de la cadena de ataque (Cyber Kill Chain), "
        "proporcionando herramientas que permiten al analista mapear la infraestructura del adversario sin interacción directa."
    )
    add_styled_paragraph(doc, taxonomia_teoria)

    add_styled_paragraph(doc, "5.3 Inteligencia Artificial Probabilística y NLP en Ciberseguridad", bold=True)
    ia_cientifica = (
        "Para la clasificación de intenciones de búsqueda, el proyecto utiliza el Clasificador Naive Bayes, un modelo basado en la probabilidad "
        "condicional del Teorema de Bayes. En el contexto de la ciberseguridad, esto se define como una forma de Procesamiento de Lenguaje Natural (NLP) "
        "aplicado a la detección de Threat Intelligence. El sistema calcula la probabilidad de que un conjunto de términos (tokens) pertenezca a una "
        "categoría de amenaza (ej. 'reputación', 'phishing') basándose en un modelo de entrenamiento supervisado, superando la limitación de los "
        "sistemas basados únicamente en firmas o expresiones regulares."
    )
    add_styled_paragraph(doc, ia_cientifica)

    add_styled_paragraph(doc, "5.4 Fundamentación Lógica y Matemática: El Teorema de Bayes", bold=True)
    bayes_logic = (
        "Desde una perspectiva lógica, el Teorema de Bayes representa un cambio de paradigma en la inferencia estadística: permite actualizar "
        "nuestra creencia o conocimiento previo (Prior) sobre un evento a medida que aparece nueva evidencia (Evidence). En el contexto de OSINT Deck, "
        "esto se traduce en la capacidad del sistema para 'aprender' que, si ciertas palabras técnicas aparecen en una consulta, la probabilidad "
        "de que el investigador esté realizando un análisis de vulnerabilidades aumenta, incluso si la consulta es ambigua. \n\n"
        "La base matemática se expresa como: P(C|D) = [P(D|C) * P(C)] / P(D). Donde: \n"
        "1) Probabilidad Prior (P(C)): Es el conocimiento inicial del sistema sobre la frecuencia de las categorías. \n"
        "2) Likelihood (P(D|C)): Es la probabilidad de observar el dato D si estuviéramos en la categoría C. \n"
        "3) Probabilidad Posterior (P(C|D)): Es el resultado final; la probabilidad actualizada de que estemos ante una intención C dado que vimos el dato D. \n\n"
        "En la implementación técnica, se utiliza 'Laplace Smoothing' (+1) para evitar que palabras desconocidas anulen el cálculo y 'Escalamiento Logarítmico' "
        "para garantizar la precisión numérica, permitiendo que el motor de decisión actúe de forma consistente y lógica ante cada nueva investigación."
    )
    add_styled_paragraph(doc, bayes_logic)

    add_styled_paragraph(doc, "5.5 Arquitectura de Capas y Seguridad por Diseño (STRIDE)", bold=True)
    arch_avanzada = (
        "La implementación de 'Clean Architecture' se fundamenta en el principio de inversión de dependencia, lo que permite una auditoría de seguridad "
        "granular. El diseño del sistema ha sido validado mediante el modelo de amenazas STRIDE (Spoofing, Tampering, Repudiation, Information Disclosure, "
        "Denial of Service, Elevation of Privilege). Esta base teórica permite que OSINT Deck no solo sea una herramienta funcional, sino un sistema "
        "resiliente que protege la integridad de la investigación y la confidencialidad de las consultas del analista."
    )
    add_styled_paragraph(doc, arch_avanzada)

    add_styled_paragraph(doc, "5.5 Marco Legal Argentino y Ética del Investigador Digital", bold=True)
    legal_avanzado = (
        "La base legal del proyecto es la Ley 25.326 de Protección de Datos Personales de la República Argentina. El marco teórico integra el "
        "concepto de 'Expectativa Razonable de Privacidad' en el ciberespacio, asegurando que la orquestación de herramientas se limite a fuentes "
        "de acceso público irrestricto. Se adopta el código de ética del profesional de ciberseguridad, implementando controles técnicos (Filtro de Toxicidad) "
        "para prevenir el uso de la herramienta en actividades de hostigamiento o ciberdelincuencia, cumpliendo con el principio de responsabilidad proactiva."
    )
    add_styled_paragraph(doc, legal_avanzado)

    # --- 6. METODOLOGÍA Y DESARROLLO ---
    add_styled_paragraph(doc, "6. Metodología, Implementación y Seguridad", level=1)
    
    add_styled_paragraph(doc, "6.1 Desarrollo de la OSINT Core Library", bold=True)
    lib_desc = (
        "El componente central del proyecto es la 'OSINT Core Library', una librería de servicios desacoplada que gestiona la detección de tipos "
        "de datos y la lógica de decisión. Su implementación bajo el estándar PSR-4 permite que la inteligencia del proyecto sea portable y "
        "auditable de forma independiente al plugin de WordPress."
    )
    add_styled_paragraph(doc, lib_desc)

    add_styled_paragraph(doc, "6.2 Auditoría de Seguridad OWASP 2025", bold=True)
    owasp_list = [
        ("A01:2025 - Control de Acceso Quebrado", "Validación estricta de privilegios administrativos en cada punto de entrada de la API."),
        ("A03:2025 - Fallas en la Cadena de Suministro", "Mitigado mediante el uso de bibliotecas locales y CDNs auditadas, evitando dependencias externas no controladas."),
        ("A05:2025 - Inyección", "Protección total mediante sentencias SQL preparadas ($wpdb->prepare) y sanitización de outputs contra XSS."),
        ("A09:2025 - Fallas en el Registro y Alerta", "Implementación de una tabla de logs persistente que permite la auditoría forense de cada acción del sistema."),
        ("A10:2025 - Mal Manejo de Condiciones Excepcionales", "Gestión robusta de errores que evita la fuga de información sensible del servidor en mensajes de error públicos.")
    ]
    for item, desc in owasp_list:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{item}: ")
        run.bold = True
        p.add_run(desc)

    # --- 7. CASO DE USO Y RESULTADOS ---
    add_styled_paragraph(doc, "7. Caso de Uso: Investigación de Filtración de Datos Corporativos", level=1)
    case_study = (
        "Se simuló un escenario de respuesta ante incidentes donde se detectó un correo electrónico corporativo en un foro de leaks. "
        "Al utilizar OSINT Deck, el motor de IA clasificó automáticamente la intención como 'Breach Analysis', desplegando herramientas de "
        "verificación de hashes y bases de datos de filtraciones. El analista logró correlacionar el email con una IP de origen y un alias de "
        "usuario en menos de 5 minutos, un proceso que manualmente requiere más de 20 minutos de navegación entre sitios externos."
    )
    add_styled_paragraph(doc, case_study)

    # --- 8. CONCLUSIÓN Y TRABAJO FUTURO ---
    add_styled_paragraph(doc, "8. Conclusión y Prospectiva", level=1)
    conclusion = (
        "OSINT Deck demuestra que es posible armonizar la simplicidad de uso con la robustez técnica y la seguridad de nivel profesional. "
        "Como trabajo futuro, se plantea la integración con modelos de lenguaje de gran escala (LLMs) para la generación automatizada de "
        "informes periciales, elevando la herramienta al nivel de un asistente de ciberinteligencia autónomo."
    )
    add_styled_paragraph(doc, conclusion)

    # --- 9. BIBLIOGRAFÍA APA 7 (EXPANDIDA) ---
    doc.add_page_break()
    add_styled_paragraph(doc, "9. Referencias Bibliográficas", level=1)
    refs = [
        "Bazzell, M. (2023). Open Source Intelligence Techniques: Resources for Digital Investigators. IntelTechniques.",
        "Clark, R. M. (2019). Intelligence Analysis: A Target-Centric Approach. CQ Press.",
        "Heuer, R. J. (1999). Psychology of Intelligence Analysis. Center for the Study of Intelligence.",
        "Kent, S. (1949). Strategic Intelligence for American World Policy. Princeton University Press.",
        "Ley de Protección de Datos Personales N° 25.326. (2000). Honorable Congreso de la Nación Argentina.",
        "OWASP Foundation. (2025). OWASP Top 10:2025 Web Application Security Risks. https://owasp.org/Top10/2025/",
        "Universidad Nacional Raúl Scalabrini Ortiz. (2019). Plan de Estudios: Licenciatura en Ciberseguridad."
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run()
        set_font(run)

    # Save
    path = r'C:\xampp\htdocs\osint.com.ar.local\wp-content\plugins\osint-deck\TFI\PTFI_OSINT_DECK_MASTER_UNSO.docx'
    doc.save(path)
    print(f"Documento Maestro blindado generado en: {path}")

if __name__ == "__main__":
    generate_master_ptfi()
