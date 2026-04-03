from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_section_header(doc, text):
    h = doc.add_heading(text, level=1)
    h.style.font.color.rgb = RGBColor(34, 113, 177) # WordPress Blue

def add_paragraph_with_updates(doc, text, update_note=None):
    p = doc.add_paragraph(text)
    if update_note:
        r = p.add_run(f"\n[ACTUALIZACIÓN TÉCNICA: {update_note}]")
        r.italic = True
        r.font.color.rgb = RGBColor(200, 0, 0)

def generate_final_tfi():
    doc = Document()
    
    # Carátula
    title = doc.add_heading('TRABAJO FINAL INTEGRADOR (TFI) - VERSIÓN PROFESIONAL v1.3.7', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Licenciatura en Ciberseguridad - UNSO').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Proyecto: OSINT Deck').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Resumen
    add_section_header(doc, 'Resumen del Proyecto')
    resumen_text = (
        "OSINT Deck es un plugin para WordPress orientado a centralizar y organizar herramientas de investigación digital. "
        "A diferencia de directorios estáticos, implementa un motor de decisión híbrido que combina expresiones regulares con "
        "Inteligencia Artificial (Clasificador Naive Bayes) para la detección automática de tipos de datos e intenciones de búsqueda. "
        "El sistema orquestará dinámicamente recursos pertinentes según la información detectada, optimizando el flujo investigativo."
    )
    add_paragraph_with_updates(doc, resumen_text, "Se integró el componente de IA para mejorar la precisión en la detección de inputs no estructurados.")

    # 2. Justificación
    add_section_header(doc, 'Justificación')
    just_text = (
        "La fragmentación de herramientas OSINT y la presencia de enlaces obsoletos en directorios tradicionales afecta la eficiencia. "
        "OSINT Deck aborda esto mediante un sistema de reporte comunitario de herramientas inactivas y una validación técnica rigurosa "
        "a través de un gestor de TLDs (Top-Level Domains) actualizado contra la base de datos de la IANA, garantizando la vigencia de los recursos."
    )
    add_paragraph_with_updates(doc, just_text, "Incorporación de TLDManager para validación técnica de dominios.")

    # 3. Objetivos
    add_section_header(doc, 'Objetivos')
    doc.add_heading('Objetivo General', level=2)
    doc.add_paragraph("Desarrollar un plugin para WordPress bajo una arquitectura de capas segura que integre y automatice el uso de herramientas OSINT.")
    
    doc.add_heading('Objetivos Específicos', level=2)
    objs = [
        "Analizar limitaciones de directorios OSINT.",
        "Implementar una Arquitectura de Capas (Domain-Driven Design) para garantizar escalabilidad y seguridad.",
        "Desarrollar un motor de IA para la clasificación dinámica de intenciones de búsqueda.",
        "Gestionar la persistencia mediante tablas SQL personalizadas para optimizar el manejo de Big Data OSINT.",
        "Auditar el sistema bajo los estándares de seguridad OWASP Top 10."
    ]
    for obj in objs:
        doc.add_paragraph(obj, style='List Bullet')

    # 4. Marco Teórico EXPANDIDO
    add_section_header(doc, 'Marco Teórico (Revisado y Expandido)')
    
    doc.add_heading('4.1 Fundamentos de Inteligencia y OSINT', level=2)
    doc.add_paragraph("Se mantienen los fundamentos de Kent (1949), Lowenthal (2022) y Heuer (1999) sobre el ciclo de inteligencia y las limitaciones cognitivas del analista, enfocando la dispersión de herramientas como un factor crítico de carga cognitiva.")
    
    doc.add_heading('4.2 Aprendizaje Automático en la Clasificación de Intenciones', level=2)
    ml_text = (
        "Para superar la rigidez de las expresiones regulares, el proyecto fundamenta su motor de búsqueda en el Clasificador Naive Bayes. "
        "Este modelo probabilístico, basado en el Teorema de Bayes, permite calcular la probabilidad de que una consulta pertenezca a una categoría "
        "específica (ej. 'reputación', 'filtración') basándose en la frecuencia de términos. Esto permite una interfaz más humana y adaptable al "
        "lenguaje natural del investigador."
    )
    add_paragraph_with_updates(doc, ml_text, "Teoría de Machine Learning aplicada a la clasificación de texto.")

    doc.add_heading('4.3 Arquitectura de Software y Seguridad por Diseño', level=2)
    arch_theory = (
        "La seguridad en el desarrollo de software (Security by Design) es un pilar de la ciberseguridad moderna. El proyecto implementa una "
        "Arquitectura de Capas que separa el Dominio (lógica de negocio), la Infraestructura (acceso a datos) y la Presentación (interfaz). "
        "Esta separación minimiza la superficie de ataque y facilita la implementación de controles de seguridad granulares, como la validación "
        "de datos en la capa de entrada y el uso de repositorios para evitar la manipulación directa de la base de datos."
    )
    add_paragraph_with_updates(doc, arch_theory, "Fundamentos de Arquitectura Limpia y Seguridad por Diseño.")

    doc.add_heading('4.4 Marco Legal y Ético de OSINT en Argentina', level=2)
    legal_text = (
        "La recolección de información de fuentes abiertas debe realizarse dentro del marco legal vigente. En Argentina, la Ley 25.326 de "
        "Protección de Datos Personales establece principios de licitud y finalidad. OSINT Deck se fundamenta en el uso ético de la información, "
        "limitando su alcance a datos de acceso público y promoviendo la transparencia mediante la clasificación profesional de los recursos utilizados."
    )
    add_paragraph_with_updates(doc, legal_text, "Marco legal argentino (Ley 25.326) y ética profesional.")

    doc.add_heading('4.5 Estándares de Seguridad Web (OWASP)', level=2)
    owasp_text = (
        "El proyecto adopta el marco de trabajo de OWASP (Open Web Application Security Project) Top 10 2021 como referencia para la mitigación "
        "de riesgos. Este estándar permite sistematizar la defensa contra las vulnerabilidades más críticas de la web, garantizando que el entorno "
        "de investigación sea resistente a ataques de inyección, fallos en el control de acceso y problemas de registro/monitoreo."
    )
    add_paragraph_with_updates(doc, owasp_text, "Metodología de desarrollo seguro basada en estándares internacionales.")

    # 5. Proceso de Desarrollo y Arquitectura
    add_section_header(doc, 'Implementación Técnica y Arquitectura')
    
    doc.add_heading('Arquitectura de Capas (Layered Architecture)', level=2)
    arch_desc = (
        "El sistema se organiza en tres capas principales:\n"
        "1. Capa de Dominio: Lógica de negocio pura (Motor de Decisión, Clasificador IA).\n"
        "2. Capa de Infraestructura: Persistencia en tablas SQL personalizadas y servicios de soporte (TLD, Logs).\n"
        "3. Capa de Presentación: Interfaz administrativa y API AJAX protegida por Nonces."
    )
    add_paragraph_with_updates(doc, arch_desc, "Implementación de Arquitectura Limpia.")

    doc.add_heading('Persistencia de Datos', level=2)
    db_text = (
        "Para garantizar el rendimiento, el plugin utiliza tablas SQL dedicadas (wp_osint_tools, wp_osint_categories) en lugar de la tabla "
        "estándar de opciones de WordPress. Esto permite realizar consultas complejas y filtrados dinámicos de forma eficiente."
    )
    add_paragraph_with_updates(doc, db_text, "Migración de JSON/Options a Tablas SQL personalizadas.")

    # 6. Auditoría de Seguridad (OWASP)
    add_section_header(doc, 'Auditoría de Seguridad y Cumplimiento')
    doc.add_paragraph("Se detallan las medidas tomadas para cumplir con el OWASP Top 10 2021:")
    
    security_points = [
        "A01:2021-Broken Access Control: Implementación de capacidades (manage_options) y Nonces en todos los endpoints.",
        "A03:2021-Injection: Uso de consultas preparadas ($wpdb->prepare) y funciones de escape (esc_html, esc_attr).",
        "A09:2021-Security Logging and Monitoring: Sistema de logs persistente con auditoría de eventos críticos.",
        "A10:2021-SSRF: Validación de URLs externas antes de la ejecución de peticiones remotas."
    ]
    for point in security_points:
        doc.add_paragraph(point, style='List Bullet')

    # 7. Guía de Instalación (Anexo)
    doc.add_page_break()
    add_section_header(doc, 'Anexo: Guía de Instalación y Despliegue')
    steps = [
        "Requisitos: WordPress 6.0+, PHP 7.4+, MySQL 5.7+.",
        "Instalación: Cargar el .zip en el panel de Plugins de WordPress.",
        "Activación: El sistema ejecutará automáticamente las migraciones SQL para crear las tablas necesarias.",
        "Puesta en marcha: Insertar el shortcode [osint_deck] en la página deseada."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # Guardar
    path = r'C:\xampp\htdocs\osint.com.ar.local\wp-content\plugins\osint-deck\TFI\TFI_Osint_Deck_Final_Revisado_v2.docx'
    doc.save(path)
    print(f"Documento final generado en: {path}")

if __name__ == "__main__":
    generate_final_tfi()
